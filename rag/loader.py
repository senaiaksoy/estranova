from __future__ import annotations

from pathlib import Path

from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LCDocument


DATA_PATH = Path("data/estranova_kaynak.docx")
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50


def load_docx_text(path: Path = DATA_PATH) -> str:
    if not path.exists():
        raise FileNotFoundError(f"Docx file not found: {path}")

    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n\n".join(paragraphs)


def build_chunks(path: Path = DATA_PATH) -> list[LCDocument]:
    text = load_docx_text(path)
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_text(text)

    return [
        LCDocument(
            page_content=chunk,
            metadata={
                "source": str(path),
                "chunk_index": index,
            },
        )
        for index, chunk in enumerate(chunks)
    ]
