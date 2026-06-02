"""
Estranova editöryal makale taslağı → Word çıktısı.

Yazara redline/yorum amacıyla göndermek için tasarlandı.
Brand renkleri (bordo/gold/mustard) ve DOCX fallback fontları
ile minimal, okunabilir bir taslak üretir.

Kullanım:
    python scripts/build-article-docx.py [astro-dosya-yolu] [output-klasör]

Default:
    astro: ../Estranova/src/pages/hormonal-gecis/40-sonrasi/saglik-kararlarinda-simdi-mi-sorusu.astro
    output: article-drafts/
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor

# Brand colors
BORDO = RGBColor(0x4f, 0x17, 0x1c)
GOLD = RGBColor(0xC9, 0xA9, 0x6E)
MUSTARD = RGBColor(0x77, 0x5A, 0x19)
CHARCOAL = RGBColor(0x2D, 0x2D, 0x2D)
GRAY = RGBColor(0x6B, 0x6B, 0x6B)

EVIDENCE_LABEL = {1: 'Sinirli', 2: 'Sinirli', 3: 'Orta', 4: 'Iyi', 5: 'Guclu'}

WRITER_NAMES = {
    'berna-aksoy': 'Berna Aksoy',
    'alara-baykent': 'Alara Baykent',
    'basak-pelister': 'Basak Pelister',
    'duygu-karaosmanoglu': 'Dt. Duygu Karaosmanoglu',
    'ozlem-denizmen': 'Ozlem Denizmen',
    'rima-erdemir': 'Rima Erdemir',
    'gamze-cizreli': 'Gamze Cizreli',
    'isik-selin-kuyumcu': 'Isik Selin Gunce',
    'demet-kizilkaya': 'Demet Kizilkaya',
    'senai-aksoy': 'Senai Aksoy',
}


def html_to_text(html):
    """Strip Astro/HTML tags; preserve **bold** and _italic_ markers; render Evidence."""
    html = re.sub(
        r'<Evidence\s+level=\{(\d+)\}\s*/>',
        lambda m: f' [Kanit: {EVIDENCE_LABEL.get(int(m.group(1)), "Orta")}]',
        html,
    )
    html = re.sub(
        r'<Evidence\s+from=\{(\d+)\}\s+to=\{(\d+)\}\s*/>',
        lambda m: f' [Kanit: {EVIDENCE_LABEL.get(int(m.group(1)))}-{EVIDENCE_LABEL.get(int(m.group(2)))}]',
        html,
    )
    html = re.sub(r'<strong[^>]*>(.+?)</strong>', r'**\1**', html, flags=re.DOTALL)
    html = re.sub(r'<em[^>]*>(.+?)</em>', r'_\1_', html, flags=re.DOTALL)
    html = re.sub(r'<br\s*/?>', ' ', html)
    html = re.sub(r'<[^>]+>', '', html)
    html = re.sub(r'\s+', ' ', html).strip()
    return html


def parse_article(text):
    title = re.search(r"const articleTitle = '(.*?)';", text).group(1)
    desc = re.search(r"const articleDescription =\s*'(.*?)';", text, re.DOTALL).group(1)
    pub = re.search(r"publishedDate: '(.*?)'", text).group(1)
    slug = re.search(r"writerSlug: '(.*?)'", text).group(1)

    summary = ''
    sm = re.search(r'Kisa\s+Ozet|Kısa\s+Özet', text)
    if sm:
        sum_match = re.search(
            r'K[ıi]sa\s+[ÖO]zet</h2>\s*<p[^>]*>(.+?)</p>',
            text,
            re.DOTALL,
        )
        if sum_match:
            summary = html_to_text(sum_match.group(1))

    prose_match = re.search(r'<ArticleProsePanel>(.+?)</ArticleProsePanel>', text, re.DOTALL)
    sections = []
    if prose_match:
        prose = prose_match.group(1)
        h2_pat = re.compile(r'<h2\s+id="([^"]+)">(.+?)</h2>(.*?)(?=<h2\s+id="|\Z)', re.DOTALL)
        for m in h2_pat.finditer(prose):
            heading = html_to_text(m.group(2))
            body = m.group(3)
            blocks = []
            for bm in re.finditer(r'(<p[^>]*>.+?</p>|<ul[^>]*>.+?</ul>)', body, re.DOTALL):
                blk = bm.group(1)
                if blk.startswith('<p'):
                    blocks.append(('p', html_to_text(blk)))
                else:
                    items = re.findall(r'<li[^>]*>(.+?)</li>', blk, re.DOTALL)
                    blocks.append(('ul', [html_to_text(it) for it in items]))
            sections.append({'heading': heading, 'blocks': blocks})

    editor = ''
    em = re.search(
        r'Bilimsel\s+Edit[öo]r\s+Notu</p>\s*<p[^>]*>(.+?)</p>',
        text,
        re.DOTALL,
    )
    if em:
        editor = html_to_text(em.group(1))

    closing = None
    cm = re.search(
        r'>\s*(Bir\s+Soru[^<]*)\s*</p>\s*<p[^>]*>(.+?)</p>\s*<p[^>]*>(.+?)</p>',
        text,
        re.DOTALL,
    )
    if cm:
        closing = {
            'label': html_to_text(cm.group(1)),
            'bridge': html_to_text(cm.group(2)),
            'question': html_to_text(cm.group(3)),
        }

    return {
        'title': title,
        'description': desc.strip(),
        'pub_date': pub,
        'writer_slug': slug,
        'summary': summary,
        'sections': sections,
        'editor_note': editor,
        'closing': closing,
    }


def add_formatted(paragraph, text, base_size=11, base_color=CHARCOAL):
    tokens = re.split(r'(\*\*.+?\*\*|_.+?_|\[Kanit:[^\]]+\])', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = paragraph.add_run(tok[2:-2])
            r.bold = True
            r.font.color.rgb = base_color
        elif tok.startswith('_') and tok.endswith('_'):
            r = paragraph.add_run(tok[1:-1])
            r.italic = True
            r.font.color.rgb = base_color
        elif tok.startswith('[Kanit:'):
            r = paragraph.add_run(tok)
            r.font.color.rgb = MUSTARD
            r.font.size = Pt(base_size - 2)
        else:
            r = paragraph.add_run(tok)
            r.font.color.rgb = base_color
        r.font.name = 'Georgia'
        if not r.font.size:
            r.font.size = Pt(base_size)


def build_docx(article, writer_name, out_path):
    doc = Document()
    s = doc.sections[0]
    s.left_margin = Inches(1.2)
    s.right_margin = Inches(1.2)
    s.top_margin = Inches(1.0)
    s.bottom_margin = Inches(1.0)

    eb = doc.add_paragraph()
    eb.paragraph_format.space_after = Pt(0)
    er = eb.add_run('ESTRANOVA - EDITORYAL TASLAK')
    er.font.name = 'Calibri'
    er.font.size = Pt(8)
    er.font.color.rgb = MUSTARD
    er.bold = True

    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(8)
    tp.paragraph_format.space_after = Pt(4)
    tr = tp.add_run(article['title'])
    tr.font.name = 'Georgia'
    tr.font.size = Pt(22)
    tr.font.color.rgb = CHARCOAL

    bp = doc.add_paragraph()
    bp.paragraph_format.space_after = Pt(20)
    br = bp.add_run(f"Yazar: {writer_name}  -  Yayin: {article['pub_date']}")
    br.font.name = 'Calibri'
    br.font.size = Pt(10)
    br.font.color.rgb = GRAY
    br.italic = True

    if article['description']:
        dp = doc.add_paragraph()
        dp.paragraph_format.space_after = Pt(24)
        dr = dp.add_run(article['description'])
        dr.font.name = 'Georgia'
        dr.font.size = Pt(12)
        dr.font.color.rgb = BORDO
        dr.italic = True

    if article['summary']:
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(12)
        hr = hp.add_run('KISA OZET')
        hr.font.name = 'Calibri'
        hr.font.size = Pt(8)
        hr.font.color.rgb = MUSTARD
        hr.bold = True

        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(20)
        add_formatted(sp, article['summary'])

    for sec in article['sections']:
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(20)
        hp.paragraph_format.space_after = Pt(8)
        hr = hp.add_run(sec['heading'])
        hr.font.name = 'Georgia'
        hr.font.size = Pt(15)
        hr.font.color.rgb = BORDO
        hr.bold = True

        for typ, content in sec['blocks']:
            if typ == 'p':
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(10)
                add_formatted(p, content)
            else:
                for item in content:
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.space_after = Pt(6)
                    add_formatted(p, item)

    if article.get('closing'):
        c = article['closing']
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(28)
        hr = hp.add_run(c['label'].upper())
        hr.font.name = 'Calibri'
        hr.font.size = Pt(8)
        hr.font.color.rgb = MUSTARD
        hr.bold = True

        bp = doc.add_paragraph()
        bp.paragraph_format.space_after = Pt(8)
        br = bp.add_run(c['bridge'])
        br.font.name = 'Georgia'
        br.font.size = Pt(13)
        br.font.color.rgb = BORDO
        br.italic = True

        qp = doc.add_paragraph()
        qp.paragraph_format.space_after = Pt(20)
        add_formatted(qp, c['question'])

    if article['editor_note']:
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(28)
        hr = hp.add_run('BILIMSEL EDITOR NOTU')
        hr.font.name = 'Calibri'
        hr.font.size = Pt(8)
        hr.font.color.rgb = BORDO
        hr.bold = True

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(20)
        add_formatted(p, article['editor_note'])

    dp = doc.add_paragraph()
    dp.paragraph_format.space_before = Pt(20)
    dr = dp.add_run(
        'Bu icerik genel bilgilendirme amaclidir ve bireysel tibbi degerlendirme, '
        'tani veya tedavinin yerini almaz. Saglik kararlari icin hekiminizle birlikte '
        'degerlendirme yapiniz.'
    )
    dr.font.name = 'Calibri'
    dr.font.size = Pt(9)
    dr.font.color.rgb = GRAY
    dr.italic = True

    doc.save(out_path)


if __name__ == '__main__':
    here = Path(__file__).resolve().parent.parent

    if len(sys.argv) > 1:
        astro_path = Path(sys.argv[1])
    else:
        astro_path = here / 'src' / 'pages' / 'hormonal-gecis' / '40-sonrasi' / 'saglik-kararlarinda-simdi-mi-sorusu.astro'

    out_dir = Path(sys.argv[2]) if len(sys.argv) > 2 else (here / 'article-drafts')
    out_dir.mkdir(parents=True, exist_ok=True)

    text = astro_path.read_text(encoding='utf-8')
    article = parse_article(text)
    writer_name = WRITER_NAMES.get(article['writer_slug'], article['writer_slug'])

    out_file = out_dir / f"{astro_path.stem}.docx"
    build_docx(article, writer_name, out_file)
    print(f'OK -> {out_file}')
